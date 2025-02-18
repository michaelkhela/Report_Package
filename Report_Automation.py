# For more details, see the Report Automation Protocol

# Script created by Michael Khela on 10/20/2023

# In[1]:
    
#Import apprpriate libraries
import re
import pandas as pd
import math
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

    # In[3]:
        
#pulls neccessary variables from the run py        
def report_fcn(root_filepath,
               sub_id,
               pls_admin,
               msel_admin,
               admin_1,
               admin_2,
               redcap_filename):
    
    # df_0 is the REDcap export CSV that contains ALL ASSESSMENTS. In order for the code to work name the file "0_REDCap_export_bridge"
    
    #msel_1 is the descriptions for the MSEL
    msel_1=pd.read_csv(root_filepath + 'Inputs/Descriptions/Mullen_Descriptions.csv')
    #df_0 (REDCap Export)
    df_0=pd.read_csv(root_filepath + 'Inputs/REDCap/' + redcap_filename, index_col ="subject_id")
    
    #pls_1 is the descriptions for the PLS
    pls_1=pd.read_csv(root_filepath + 'Inputs/Descriptions/PLS_Descriptions.csv')
    
    ########################## CHANGE FOR EACH STUDY ################################################################
    #change redcap_event_name, visit #, and the childname for the report for each study
        
    #separates the visits to their respective ones based on the answer to the visit input
    df_visit = df_0[df_0['redcap_event_name'].str.contains('visit_1')]
    sub_id = int(sub_id)
    df_sub_visit = df_visit.loc[sub_id].reset_index().rename(columns={'index': 'Var'})
    
    #used to create the filename (last, f.)
    child_name_series = df_0.loc[sub_id, 'child_full_name'].dropna().str.split()
    if not child_name_series.empty:
        first_name = child_name_series.str[0].iloc[0]
        last_name = child_name_series.str[-1].iloc[0]
        formatted_name = f"{last_name}, {first_name[0].upper()}."
    
    #################################################################################################################
    # In[4]:
        
    
    # MSEL Table Creation
    
    # Data extracton for MSEL (t-scores, pr, ae)
    
    def extract_value(df_sub_visit, var_prefix, sub_id):
        var_t = df_sub_visit[df_sub_visit['Var'].str.contains(f'{var_prefix}_t')][sub_id].iloc[0]
        var_pr = df_sub_visit[df_sub_visit['Var'].str.contains(f'{var_prefix}_pr')][sub_id].iloc[0]
        var_ae = df_sub_visit[df_sub_visit['Var'].str.contains(f'{var_prefix}_ae')][sub_id].iloc[0]
        return var_t, var_pr, var_ae
    
    #extract values for visual_reception (vr: t_scores, percentile rank, age equivalence)
    vr_t, vr_pr, vr_ae = extract_value(df_sub_visit, 'visual_reception', sub_id)
    
    #extract values for fine_motor (fm: t_scores, percentile rank, age equivalence)
    fm_t, fm_pr, fm_ae = extract_value(df_sub_visit, 'fine_motor', sub_id)
    
    # In[5]:
        
    # MSEL Table Insertion 
    
    #creates the table that will hold the t, ae, pr scores
    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)
    
    ################################# CHANGE FILE PATH FOR EACH STUDY ########################################
    doc = Document(root_filepath +'Inputs/Templates/0_Report_Template.docx')
    
    #finds the text "insert_msel_table" and replaces it with the msel table
    for paragraph in doc.paragraphs:
        if "[insert_msel_table]" in paragraph.text:
            #clear the existing content ([insert_msel_table])
            paragraph.clear()
            
            #create a 3x4 table
            table = doc.add_table(rows=3, cols=4)
            table.style = 'Table Grid'
            
            #add a style to the table
            table.autofit = True
            table.allow_autofit = True
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            table.cell(0, 0).text = "Mullen Scales of Early Learning"  
            table.cell(0, 1).text = "T Scores"  
            table.cell(0, 2).text = "Percentile Rank"  
            table.cell(0, 3).text = "Age Equivalent"  
            table.cell(1, 0).text = "Visual Reception"  
            table.cell(2, 0).text = "Fine Motor" 
            
            ##### CUSTOM CHANGE TO ACCODOMATE VALUE FLIPPING
            # Convert vr_ae to months
            try:
                vr_months = int(vr_ae)
            except ValueError:
                vr_months = str(vr_ae)
            
            # Convert fm_ae to months
            try:
                fm_months = int(fm_ae)
            except ValueError:
                fm_months = str(fm_ae)
            
            #formats the float values to have two decimal places and add " months" at the end
            table.cell(1, 1).text = "{:.0f}".format(vr_t)  
            table.cell(1, 2).text = "{:.0f}".format(vr_pr)  
            table.cell(1, 3).text = f"{(vr_months)} months" 
            table.cell(2, 1).text = "{:.0f}".format(fm_t)  
            table.cell(2, 2).text = "{:.0f}".format(fm_pr)  
            table.cell(2, 3).text = f"{(fm_months)} months" 
    
            for row_num, row in enumerate(table.rows):
                for col_num, cell in enumerate(row.cells):
                    cell.width = Inches(2.7) 
                    cell.height = Inches(0.4)
                    for cell_paragraph in cell.paragraphs:
                        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if col_num == 0:
                            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for run in cell_paragraph.runs:
                            run.font.name = 'Calibri Light'
                            run.font.size = Pt(12)
                            run.bold = False
    
            #function to apply formatting to the entire first row of the table
            def format_first_row(table, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
                first_row = table.rows[0]
                for cell in first_row.cells:
                    for cell_paragraph in cell.paragraphs:
                        cell_paragraph.alignment = alignment
                        for run in cell_paragraph.runs:
                            run.bold = True
                            run.font.name = 'Calibri Light'
                            run.font.size = Pt(12)
    
            #format the entire first row 
            format_first_row(table)
            
            #use the move_table_after function to insert the table after the paragraph
            move_table_after(table, paragraph)
    
    
    # In[6]:
     
    # VABS Table Creation (Domain and Subdomain)
      
    # Data extracton for VABS (ss, pr, vscale, ae)
    
    def extract_value(df_sub_visit, var_prefix, sub_id, suffix=''):
        var_name = f'{var_prefix}_{suffix}'.rstrip('_')
        var_value = df_sub_visit[df_sub_visit['Var'].str.contains(f'{var_name}')][sub_id].iloc[0]
        #removes '<' symbol if present and the variable name contains 'pr'
        if 'pr' in var_name.lower() and isinstance(var_value, str):
            var_value = var_value.replace('<', '')
        return var_value
    
    ########################## CHANGE FOR EACH STUDY ################################################################
    #extract values for ss
    vabs_abc_ss = extract_value(df_sub_visit, 'vas_abc', sub_id, 'ss')
    vabs_comm_ss = extract_value(df_sub_visit, 'vas_communication', sub_id, 'ss')
    vabs_dl_ss = extract_value(df_sub_visit, 'vas_dailyliving', sub_id, 'ss')
    vabs_social_ss = extract_value(df_sub_visit, 'vas_social', sub_id, 'ss')
    vabs_motor_ss = extract_value(df_sub_visit, 'vas_motor', sub_id, 'ss')
    
    #extract values for pr
    vabs_abc_pr = extract_value(df_sub_visit, 'vas_abc', sub_id, 'pr')
    vabs_comm_pr = extract_value(df_sub_visit, 'vas_communication', sub_id, 'pr')
    vabs_dl_pr = extract_value(df_sub_visit, 'vas_dailyliving', sub_id, 'pr')
    vabs_social_pr = extract_value(df_sub_visit, 'vas_social', sub_id, 'pr')
    vabs_motor_pr = extract_value(df_sub_visit, 'vas_motor', sub_id, 'pr')
    
    #extract values for vscale
    vabs_receptive_vscale = extract_value(df_sub_visit, 'vas_receptive', sub_id, 'vscale')
    vabs_expressive_vscale = extract_value(df_sub_visit, 'vas_expressive', sub_id, 'vscale')
    vabs_written_vscale = extract_value(df_sub_visit, 'vas_written', sub_id, 'vscale')
    vabs_personal_vscale = extract_value(df_sub_visit, 'vas_personal', sub_id, 'vscale')
    vabs_domestic_vscale = extract_value(df_sub_visit, 'vas_domestic', sub_id, 'vscale')
    vabs_community_vscale = extract_value(df_sub_visit, 'vas_community', sub_id, 'vscale')
    vabs_relationship_vscale = extract_value(df_sub_visit, 'vas_interpersonal', sub_id, 'vscale')
    vabs_play_vscale = extract_value(df_sub_visit, 'vas_play', sub_id, 'vscale')
    vabs_coping_vscale = extract_value(df_sub_visit, 'vas_coping', sub_id, 'vscale')
    vabs_gross_vscale = extract_value(df_sub_visit, 'vas_gross_motor', sub_id, 'vscale')
    vabs_fine_vscale = extract_value(df_sub_visit, 'vas_fine_motor', sub_id, 'vscale')
    
    #extract values for ae
    vabs_receptive_ae = extract_value(df_sub_visit, 'vas_receptive', sub_id, 'age')
    vabs_expressive_ae = extract_value(df_sub_visit, 'vas_expressive', sub_id, 'age')
    vabs_written_ae = extract_value(df_sub_visit, 'vas_written', sub_id, 'age')
    vabs_personal_ae = extract_value(df_sub_visit, 'vas_personal', sub_id, 'age')
    vabs_domestic_ae = extract_value(df_sub_visit, 'vas_domestic', sub_id, 'age')
    vabs_community_ae = extract_value(df_sub_visit, 'vas_community', sub_id, 'age')
    vabs_relationship_ae = extract_value(df_sub_visit, 'vas_interpersonal', sub_id, 'age')
    vabs_play_ae = extract_value(df_sub_visit, 'vas_play', sub_id, 'age')
    vabs_coping_ae = extract_value(df_sub_visit, 'vas_coping', sub_id, 'age')
    vabs_gross_ae = extract_value(df_sub_visit, 'vas_gross_motor', sub_id, 'age')
    vabs_fine_ae = extract_value(df_sub_visit, 'vas_fine_motor', sub_id, 'age')   
    ############################################################################################################
    
    # In[7]:
    
    # VABS Table Insertion (Domain Related)
      
    #function to move a table after a paragraph
    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._element
        p.addnext(tbl)
        p.getparent().remove(p)
        
    #finds the text "[insert_vabs_table_1]" and replaces it with the 1st vabs table
    for paragraph in doc.paragraphs:
        if "[insert_vabs_table_1]" in paragraph.text:
            #clear the existing content
            for run in paragraph.runs:
                run.clear()
            
            #create a 6x3 table
            table = doc.add_table(rows=6, cols=3)
            table.style = 'Table Grid'
    
            #header row
            table.cell(0, 0).text = "Domain"
            table.cell(0, 1).text = "Standard Score"
            table.cell(0, 2).text = "Percentile Rank"
    
            #data rows
            data_rows = [
                ("Adaptive Behavior Composite", vabs_abc_ss, vabs_abc_pr),
                ("Communication", vabs_comm_ss, vabs_comm_pr),
                ("Daily Living Skills", vabs_dl_ss, vabs_dl_pr),
                ("Socialization", vabs_social_ss, vabs_social_pr),
                ("Motor Skills", vabs_motor_ss, vabs_motor_pr),
            ]
    
            for row, (domain, ss, pr) in enumerate(data_rows, start=1):
                table.cell(row, 0).text = domain
                table.cell(row, 1).text = "{:.0f}".format(ss)
                table.cell(row, 2).text = str(pr)
    
            for row_num, row in enumerate(table.rows):
                for col_num, cell in enumerate(row.cells):
                    cell.width = Inches(2.5) 
                    cell.height = Inches(0.4)
                    for cell_paragraph in cell.paragraphs:
                        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if col_num == 0:
                            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for run in cell_paragraph.runs:
                            run.font.name = 'Calibri Light'
                            run.font.size = Pt(12)
                            run.bold = False
    
            #function to apply formatting to the entire first row of the table
            def format_first_row(table, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
                first_row = table.rows[0]
                for cell in first_row.cells:
                    for cell_paragraph in cell.paragraphs:
                        cell_paragraph.alignment = alignment
                        for run in cell_paragraph.runs:
                            run.bold = True
                            run.font.name = 'Calibri Light'
                            run.font.size = Pt(12)
    
            #format the entire first row
            format_first_row(table)
            
            #use the move_table_after function to insert the table after the paragraph
            move_table_after(table, paragraph)
    
    
    # ### VABS Table Insertion (Subdomain Related)
    
    # In[8]:
     
    # VABS Table Insertion (Subdomain Related)  
    
    #function to move a table after a paragraph
    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._element
        p.addnext(tbl)
        p.getparent().remove(p)
        
    #data rows
    data_rows = [
        ("Communication", "", ""),
        ("    Receptive", vabs_receptive_vscale, vabs_receptive_ae),
        ("    Expressive", vabs_expressive_vscale, vabs_expressive_ae),
        ("    Written", vabs_written_vscale, vabs_written_ae),
        ("Daily Living Skills", "", "", ""),
        ("    Personal", vabs_personal_vscale, vabs_personal_ae),
        ("    Domestic", vabs_domestic_vscale, vabs_domestic_ae),
        ("    Community", vabs_community_vscale, vabs_community_ae),
        ("Socialization", "", "", ""),
        ("    Interpersonal Relationships", vabs_relationship_vscale, vabs_relationship_ae),
        ("    Play and Leisure", vabs_play_vscale, vabs_play_ae),
        ("    Coping Skills", vabs_coping_vscale, vabs_coping_ae),
        ("Motor Skills", "", ""),
        ("    Gross Motor", vabs_gross_vscale, vabs_gross_ae),
        ("    Fine Motor", vabs_fine_vscale, vabs_fine_ae)
    ]
    
    #find and replace the "[insert_vabs_table_2]" placeholder
    for paragraph in doc.paragraphs:
        if "[insert_vabs_table_2]" in paragraph.text:
            #clear the existing content
            for run in paragraph.runs:
                run.clear()
                
            #create a 16x3 table
            table = doc.add_table(rows=len(data_rows) + 1, cols=3)
            table.style = 'Table Grid'
            
            #merges rows 
            merge_rows = [1, 5, 9, 13]
            for merge_row in merge_rows:
                table.cell(merge_row, 0).merge(table.cell(merge_row, 2))
    
            #add data to the table
            for row, data_tuple in enumerate(data_rows, start=1):
                table.cell(row, 0).text = data_tuple[0]  # domain
                try:
                    table.cell(row, 1).text = "{:.0f}".format(data_tuple[1]) if len(data_tuple) > 1 and data_tuple[1] is not None else ""
                except (TypeError, ValueError):
                    table.cell(row, 1).text = ""
                try:
                    formatted_ae = str(data_tuple[2]) if len(data_tuple) > 2 and data_tuple[2] is not None else ""
                    table.cell(row, 2).text = formatted_ae
                except (TypeError, ValueError):
                    table.cell(row, 2).text = ""
                    formatted_ae = ""
                    
            #merged row naming
            table.cell(1, 0).text = "Communication"
            table.cell(5, 0).text = "Daily Living Skills"
            table.cell(9, 0).text = "Socialization"
            table.cell(13, 0).text = "Motor Skills"
                
            for row_num, row in enumerate(table.rows):
                for col_num, cell in enumerate(row.cells):
                    cell.width = Inches(2.5) 
                    cell.height = Inches(0.4)
                    for cell_paragraph in cell.paragraphs:
                        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if col_num == 0:
                            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for run in cell_paragraph.runs:
                            run.font.name = 'Calibri Light'
                            run.font.size = Pt(12)
                            run.bold = False
           
            #set the headers and fix the alignment of the merged cells
            headers = ["Subdomain", "V-Scale Score", "Age Equivalent (Years:Months)"]
            for col, header in enumerate(headers):
                table.cell(0, col).text = header
                table.cell(0, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.cell(1, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                table.cell(5, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                table.cell(9, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                table.cell(13, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in table.cell(0, col).paragraphs[0].runs:
                    run.font.name = 'Calibri Light'
                    run.font.size = Pt(12)
                    run.bold = True
    
            #use the move_table_after function to insert the table after the paragraph
            move_table_after(table, paragraph)
            
    
    # In[9]:
    
    # PLS Table Creation
    
    # Data extracton for PLS (ss, pr, ae, and correct/incorrect values)
    def extract_value(df_sub_visit, var_prefix, sub_id, suffix=''):
        var_name = f'{var_prefix}_{suffix}'.rstrip('_')
        try:
            var_value = df_sub_visit[df_sub_visit['Var'].str.contains(f'{var_name}')][sub_id].iloc[0]
            return var_value
        except (IndexError, KeyError):
            return None  #return None if the value is not found
    
    def extract_int_value(df_sub_visit, var_prefix, sub_id, suffix=''):
        var_value = extract_value(df_sub_visit, var_prefix, sub_id, suffix)
        try:
            return int(var_value) if var_value is not None else None
        except (ValueError, TypeError):
            return None  #return None if the value cannot be converted to an integer
    
    ########################## CHANGE FOR EACH STUDY ################################################################
    
    #extract values for ss
    pls_aud_ss = extract_int_value(df_sub_visit, 'pls_aud_comp', sub_id, 'ss')
    pls_exp_ss = extract_int_value(df_sub_visit, 'pls_exp_comm', sub_id, 'ss')
    pls_total_ss = extract_int_value(df_sub_visit, 'pls_total', sub_id, 'ss_2')
    
    #extract values for pr
    pls_aud_pr = extract_int_value(df_sub_visit, 'pls_aud_comp', sub_id, 'pr')
    pls_exp_pr = extract_int_value(df_sub_visit, 'pls_exp_comm', sub_id, 'pr')
    pls_total_pr = extract_int_value(df_sub_visit, 'pls_total', sub_id, 'pr')
    
    #extract values for ae
    pls_aud_ae = extract_value(df_sub_visit, 'pls_aud_comp', sub_id, 'ae_ym')
    pls_exp_ae = extract_value(df_sub_visit, 'pls_exp_comm', sub_id, 'ae_ym')
    pls_total_ae = extract_value(df_sub_visit, 'pls_total', sub_id, 'ae_ym')
    
    #extract values for correct values
    ec_corr_1 = extract_int_value(df_sub_visit, 'pls_ec_correct', sub_id, '1')
    ec_corr_2 = extract_int_value(df_sub_visit, 'pls_ec_correct', sub_id, '2')
    ec_corr_3 = extract_int_value(df_sub_visit, 'pls_ec_correct', sub_id, '3')
    ec_corr_4 = extract_int_value(df_sub_visit, 'pls_ec_correct', sub_id, '4')
    ec_corr_5 = extract_int_value(df_sub_visit, 'pls_ec_correct', sub_id, '5')
    
    #extract values for incorrect values
    ec_zeros_1 = extract_int_value(df_sub_visit, 'pls_ec_incorrect', sub_id, '1')
    ec_zeros_2 = extract_int_value(df_sub_visit, 'pls_ec_incorrect', sub_id, '2')
    ec_zeros_3 = extract_int_value(df_sub_visit, 'pls_ec_incorrect', sub_id, '3')
    ec_zeros_4 = extract_int_value(df_sub_visit, 'pls_ec_incorrect', sub_id, '4')
    ec_zeros_5 = extract_int_value(df_sub_visit, 'pls_ec_incorrect', sub_id, '5')
    
    #extract values for correct values
    ac_corr_1 = extract_int_value(df_sub_visit, 'pls_ac_correct', sub_id, '1')
    ac_corr_2 = extract_int_value(df_sub_visit, 'pls_ac_correct', sub_id, '2')
    ac_corr_3 = extract_int_value(df_sub_visit, 'pls_ac_correct', sub_id, '3')
    ac_corr_4 = extract_int_value(df_sub_visit, 'pls_ac_correct', sub_id, '4')
    ac_corr_5 = extract_int_value(df_sub_visit, 'pls_ac_correct', sub_id, '5')
    
    #extract values for incorrect values
    ac_zeros_1 = extract_int_value(df_sub_visit, 'pls_ac_incorrect', sub_id, '1')
    ac_zeros_2 = extract_int_value(df_sub_visit, 'pls_ac_incorrect', sub_id, '2')
    ac_zeros_3 = extract_int_value(df_sub_visit, 'pls_ac_incorrect', sub_id, '3')
    ac_zeros_4 = extract_int_value(df_sub_visit, 'pls_ac_incorrect', sub_id, '4')
    ac_zeros_5 = extract_int_value(df_sub_visit, 'pls_ac_incorrect', sub_id, '5')
    
    ###################################################################################################################
    
    #reformatting the ae to accommadate the y:m (12/9/23)
    def format_ae_value(ae_value):
        ae_value = ae_value.strip()
        
        if ':' in ae_value:
            # for values already in the format 1:3
            return ae_value
        elif 'y' in ae_value and 'm' in ae_value:
            # convert format 3y5m to 3:5
            match = re.match(r'(\d+)y(\d+)m', ae_value)
            if match:
                years, months = map(int, match.groups())
                return f'{years}:{months}'
            else:
                raise ValueError(f"Unsupported format for ae_value: {ae_value}")
        else:
            # replace "." with ":"
            return ae_value.replace('.', ':')
    
    
    #formatted variables:
    pls_aud_ae_formatted = format_ae_value(pls_aud_ae)
    pls_exp_ae_formatted = format_ae_value(pls_exp_ae)
    pls_total_ae_formatted = format_ae_value(pls_total_ae)
    
    
    # In[10]:
            
    # PLS Table Insertion
    
    #creates the table that will hold the ss, ae, pr scores
    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)
    
    #finds the text "insert_pls_table" and replaces it with the pls table
    for paragraph in doc.paragraphs:
        if "[insert_pls_table]" in paragraph.text:
            #clear the existing content
            paragraph.clear()
            
            #create a 4x4 table
            table = doc.add_table(rows=4, cols=4)
            table.style = 'Table Grid'
            
            table.cell(0, 0).text = "Preschool Language Scale"  
            table.cell(0, 1).text = "Standard Scores" 
            table.cell(0, 2).text = "Percentile Rank"  
            table.cell(0, 3).text = "Age Equivalent (years:months)"  
            table.cell(1, 0).text = "Auditory Comprehension"  
            table.cell(2, 0).text = "Expressive Communication" 
            table.cell(3, 0).text = "Total Language"  
    
            # Format the float values to have two decimal places and add " months" at the end
            # Function to format the float values
            def format_float_value(value):
                return "{:.0f}".format(value) if value is not None else ""
    
            # Function to format the ae values
            def format_ae_value(ae_value):
                return ae_value if ae_value is not None else ""
    
            # Set the cell values with formatting
            table.cell(1, 1).text = format_float_value(pls_aud_ss)
            table.cell(1, 2).text = format_float_value(pls_aud_pr)
            table.cell(1, 3).text = format_ae_value(pls_aud_ae_formatted)
            table.cell(2, 1).text = format_float_value(pls_exp_ss)
            table.cell(2, 2).text = format_float_value(pls_exp_pr)
            table.cell(2, 3).text = format_ae_value(pls_exp_ae_formatted)
            table.cell(3, 1).text = format_float_value(pls_total_ss)
            table.cell(3, 2).text = format_float_value(pls_total_pr)
            table.cell(3, 3).text = format_ae_value(pls_total_ae_formatted)
    
            #set cell margins to create a border around the table
            for row_num, row in enumerate(table.rows):
                for col_num, cell in enumerate(row.cells):
                    cell.width = Inches(2.7) 
                    cell.height = Inches(0.4)
                    for cell_paragraph in cell.paragraphs:
                        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if col_num == 0:
                            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for run in cell_paragraph.runs:
                            run.font.name = 'Calibri Light'
                            run.font.size = Pt(12)
                            run.bold = False
    
            #function to apply formatting to the entire first row of the table
            def format_first_row(table, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
                first_row = table.rows[0]
                for cell in first_row.cells:
                    for cell_paragraph in cell.paragraphs:
                        cell_paragraph.alignment = alignment
                        for run in cell_paragraph.runs:
                            run.bold = True
                            run.font.name = 'Calibri Light'
                            run.font.size = Pt(12)
                            
            #format the entire first row with bold text
            format_first_row(table)
    
            #use the move_table_after function to insert the table after the paragraph
            move_table_after(table, paragraph)
    
    ########################## CHANGE FOR EACH STUDY ################################################################        
    #specify the directory where you want to save the file
    save_directory = root_filepath + 'Created/'
    
    #save the modified document with the desired filename and directory
    doc.save(save_directory + formatted_name + ' Visit Report.docx')
    
    
    # In[11]:
     
    # Report Finalization
    
    # Information Related   
    
    #extracts the child's first name
    child_name = df_0.loc[sub_id, 'child_full_name'].dropna().str.split().str[0].iloc[0]
    
    #extracts the child's full name
    child_full_name = df_0.loc[sub_id, 'child_full_name'].dropna().iloc[0]
    
    #extracts the date of eval
    date_eval = pd.to_datetime(df_visit.loc[sub_id, 'visit_date'])
    date_eval_str = date_eval.strftime('%m-%d-%Y')
    
    #extracts dob then converts to str
    dob = pd.to_datetime(df_0.loc[sub_id, 'dob'].dropna().iloc[0])
    dob_str = dob.strftime('%m-%d-%Y') 
    
    # Given age in months
    age_months_decimal = df_visit.loc[sub_id, 'age_at_vist']
    
    # Extract years and remaining months, rounding down
    age_years = math.floor(age_months_decimal // 12)
    remaining_months = math.floor(age_months_decimal % 12)
    
    # Combine age in years and remaining months into a string
    age_at_eval = f"{age_years} years, {remaining_months} months"
    
    ########################## CHANGE FOR EACH STUDY ################################################################
    #adds a description for the specificed admins
    if pls_admin == '1':
        admin_pls = "Michael Khela"
    elif pls_admin == '2':
        admin_pls = "Michael Khela"
    elif pls_admin == '3':
        admin_pls = "Michael Khela"
    else:
        admin_pls = "Invalid input. Please enter 1, 2, or 3 for the corresponding administrator."
        
    if msel_admin == '1':
        admin_msel = "Michael Khela"
    elif msel_admin == '2':
        admin_msel = "Michael Khela"
    elif msel_admin == '3':
        admin_msel = "Michael Khela"
    else:
        admin_msel = "Invalid input. Please enter 1, 2, or 3 for the corresponding administrator."
        
    #adds names of admins who were at the visit
    if admin_1 == '1':
        first_examiner = "Michael Khela, Tufts University Student"
    elif admin_1 == '2':
        first_examiner = "Michael Khela, Tufts University Student"
    elif admin_1 == '3':
        first_examiner = "Michael Khela, Tufts University Student"
    else:
        first_examiner = "Invalid input. Please enter 1, 2, or 3 for the corresponding administrator."
        
    if admin_2 == '1':
        second_examiner = "Michael Khela, Wilkinson Lab Student"
    elif admin_2 == '2':
        second_examiner = "Michael Khela, Wilkinson Lab Student"
    elif admin_2 == '3':
        second_examiner = "Michael Khela, Wilkinson Lab Student"
    else:
        second_examiner = "Invalid input. Please enter 1, 2, or 3 for the corresponding administrator."
        
    #if DS (only used for kiddos with DS) (BRIDGE SPECIFIC)
    if any(df_sub_visit.apply(lambda row: 'arm_4' in str(row.values), axis=1)):
        ds_pi = "Michael Khela, in several years will be MD"
    else:
        ds_pi = ""
    
    #msel age (used for year and month in DOCX)
    if vr_ae > fm_ae:
        msel_larger_age = vr_ae
        msel_smaller_age = fm_ae
    else:
        msel_larger_age = fm_ae
        msel_smaller_age = vr_ae
        
    #PLS age (used for year and month in DOCX)
    def process_age(age):
        years_result, months_result = None, None
        
        if isinstance(age, (float, int)):
            years_result = int(age)
            months_result = int((age - years_result) * 12)
        elif isinstance(age, str):
            try:
                if '.' in age:
                    # Handle ages in the format "1.5"
                    years_result, months_result = map(int, age.split('.'))
                else:
                    # Handle ages in the format "1:5"
                    years_result, months_result = map(int, age.split(':'))
            except ValueError:
                pass
        
        return years_result, months_result
    
    #process each variable and get separate output variables
    years_aud, months_aud = process_age(pls_aud_ae_formatted)
    years_exp, months_exp = process_age(pls_exp_ae_formatted)
    years_total, months_total = process_age(pls_total_ae_formatted)
    
    # In[12]:
        
    # MSEL Description
        
    merged_sub = pd.merge(df_sub_visit, msel_1, on='Var', how='right')
    merged_sub = merged_sub.dropna()
    vr = merged_sub[merged_sub['Var'].str.contains('vr_')]
    fm = merged_sub[merged_sub['Var'].str.contains('fm_')]
    
    def extract_zeros(df_sub_visit, sub_id, condition, num_items):
        items = []
        des_items = []
    
        for _, row in df_sub_visit.iterrows():
            var = row['Var']
            value = row[sub_id]
            des = row["Des"]
    
            if condition(value):
                items.append((var, value, des))
                des_items.append(des)
    
                if len(items) > num_items:
                    items.pop(0)
                    des_items.pop(0)
    
                if len(items) == num_items:
                    break
    
        return items, des_items
    
    def extract_correct_reverse(df_sub_visit, sub_id, condition, num_items):
        items = []
        des_items = []
    
        for _, row in df_sub_visit.iloc[::-1].iterrows():
            var = row['Var']
            value = row[sub_id]
            des = row["Des"]
    
            if condition(value):
                items.append((var, value, des))
                des_items.append(des)
    
                if len(items) > num_items:
                    items.pop(0)
                    des_items.pop(0)
    
                if len(items) == num_items:
                    break
    
        return items, des_items
    
    #find and extract the last three correct items in 'vr' and pull their respective description into a list (in reverse order)
    three_correct_vr, des_correct_vr = extract_correct_reverse(vr, sub_id, lambda value: value > 0, 3)
    
    #find and extract the last three correct items in 'fm' and pull their respective description into a list (in reverse order)
    three_correct_fm, des_correct_fm = extract_correct_reverse(fm, sub_id, lambda value: value > 0, 3)
    
    #find and extract the first three zeros in 'vr' and pull their respective description into a list (in regular order)
    zeros_vr, des_zeros_vr = extract_zeros(vr, sub_id, lambda value: value == 0, 3)
    
    #find and extract the first three zeros in 'fm' and pull their respective description into a list (in regular order)
    zeros_fm, des_zeros_fm = extract_zeros(fm, sub_id, lambda value: value == 0, 3)
    
    
    # In[13]:
            
    # PLS Descriptions
    
    #assuming your original variables are not lists
    ec_corr_values = [ec_corr_1, ec_corr_2, ec_corr_3, ec_corr_4, ec_corr_5]
    ac_corr_values = [ac_corr_1, ac_corr_2, ac_corr_3, ac_corr_4, ac_corr_5]
    ec_zeros_values = [ec_zeros_1, ec_zeros_2, ec_zeros_3, ec_zeros_4, ec_zeros_5]
    ac_zeros_values = [ac_zeros_1, ac_zeros_2, ac_zeros_3, ac_zeros_4, ac_zeros_5]
    
    def extract_des(pls_1, var_suffix, indices):
        descriptions = []
    
        for idx in indices:
            ac_str = f'{var_suffix}_{idx}'
            matching_rows = pls_1[pls_1['PLS'].str.contains(ac_str)]
    
            if not matching_rows.empty:
                description = matching_rows['Item Description'].iloc[0]
                descriptions.append(description)
            else:
                descriptions.append(f'No description found for {ac_str}')
    
        return descriptions
    
    ec_corr_descriptions = extract_des(pls_1, 'EC', ec_corr_values)
    ac_corr_descriptions = extract_des(pls_1, 'AC', ac_corr_values)
    ec_zeros_descriptions = extract_des(pls_1, 'EC', ec_zeros_values)
    ac_zeros_descriptions = extract_des(pls_1, 'AC', ac_zeros_values)
    
    
    # In[14]:
        
    # Adding the Report Info + Descriptions to DOCX
    
    ########################## CHANGE FOR EACH STUDY ################################################################
    text_mappings = {
        "Childname": child_name,
        "firstlastname": child_full_name,
        "dob": dob_str,
        "age_eval": age_at_eval,
        "eval_date": date_eval_str,
        "examiner1": first_examiner,
        "examiner2": second_examiner,
        "mseladmin": admin_msel,
        "plsadmin": admin_pls,
        "dspi": ds_pi,
        "mselmonth1": str(msel_smaller_age),
        "mselmonth2": str(msel_larger_age),
        "yraud": str(years_aud),
        "mthaud": str(months_aud),
        "yrexp": str(years_exp),
        "mthexp": str(months_exp),
        "totalyear": str(years_total),
        "mthtotal": str(months_total),
        "vrcorr": des_correct_vr,
        "vrzeros": des_zeros_vr,
        "fmcorr": des_correct_fm,
        "fmzeros": des_zeros_fm,
        "accorr": ac_corr_descriptions,
        "incorrac": ac_zeros_descriptions,
        "eccorr": ec_corr_descriptions,
        "zerosec": ec_zeros_descriptions
    }
    
    doc = Document(save_directory + formatted_name + ' Visit Report.docx')
    
    #iterate through all paragraphs in the document
    for paragraph in doc.paragraphs:
        for old_text, new_text in text_mappings.items():
            if old_text in paragraph.text:
                #iterate through runs and find the run containing the old text
                for run in paragraph.runs:    
                    if old_text in run.text:
                        run.bold = True
                        #if the new text is a list, join the elements into a single string
                        if isinstance(new_text, list):
                            new_text = ", ".join(map(str, new_text))
                        #replace the old text in the run with the new text
                        run.text = run.text.replace(old_text, str(new_text))
    
                #adjust paragraph formatting as needed
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    #save the modified document with the desired filename and directory
    doc.save(save_directory + formatted_name + ' Visit Report.docx')
    
    
    # ## Inserting Signatures
    
    # In[15]:
    
    # Inserting Signatures
    #define file paths for different values of pls_admin and msel_admin
    
    ########################## CHANGE FOR EACH STUDY ################################################################
    admin_paths = {
        1: {
            '1': root_filepath + "Inputs/Signatures/mk_signature.png",
            '2': root_filepath + "Inputs/Signatures/mk_signature.png",
        },
        2: {
            '1': root_filepath + "Inputs/Signatures/mk_signature.png",
            '2': root_filepath + "Inputs/Signatures/mk_signature.png",
        },
        3: {
            '1': root_filepath + "Inputs/Signatures/mk_signature.png",
            '2': root_filepath + "Inputs/Signatures/mk_signature.png",
        },
    }
    
    #load the document
    doc = Document(save_directory + formatted_name + ' Visit Report.docx')
    
    #search for the text to replace in the entire document
    for paragraph in doc.paragraphs:
        if "[insert_signatures]" in paragraph.text:
            #clear the existing content of the paragraph
            for run in paragraph.runs:
                run.clear()
    
            #add a new run to the paragraph
            run = paragraph.add_run()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
            #fetch the corresponding image paths based on admin_[] value
            admin1_sig = int(admin_1)
            admin2_sig = int(admin_2)
    
            if admin1_sig in admin_paths and '1' in admin_paths[admin1_sig]:
                admin1_image_path = admin_paths[admin1_sig]['1']
            else:
                admin1_image_path = ''
    
            if admin2_sig in admin_paths and '2' in admin_paths[admin2_sig]:
                admin2_image_path = admin_paths[admin2_sig]['2']
            else:
                admin2_image_path = ''
    
            #insert the corresponding images based on admin_1 and admin_2 values
            run.add_picture(admin1_image_path, width=Inches(2.22))
            
            #add a new run to the paragraph
            run = paragraph.add_run()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            run.add_picture(admin2_image_path, width=Inches(2.22))
    
    ########################## CHANGE FOR EACH STUDY ################################################################
    #save the modified document
    doc.save(save_directory + formatted_name + ' Visit Report.docx')
    
    #print statement indicating that the code has been executed
    print("Congratulations!! Report for the kiddo has been created!!!")
    # In[ ]: